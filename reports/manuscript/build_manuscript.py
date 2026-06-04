from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
EVAL_DIR = ROOT / "reports" / "lea_evaluation"
OUT_DIR = ROOT / "reports" / "manuscript"
DOCX_PATH = OUT_DIR / "qmof_lea_manuscript_draft.docx"
ARCH_FIGURE = ROOT / "manuscript" / "figures" / "qmof_ai_platform_architecture.png"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"
MID_GRAY = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            set_cell_width(row.cells[idx], width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles.add_style("CaptionCustom", 1)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MID_GRAY
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.05

    body_small = styles.add_style("BodySmall", 1)
    body_small.font.name = "Calibri"
    body_small.font.size = Pt(9)
    body_small.paragraph_format.space_after = Pt(4)
    body_small.paragraph_format.line_spacing = 1.05

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("QMOF-LEA manuscript draft | Page ")
    add_field(footer, "PAGE")


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(
        "Lotus Effect Optimization-Guided Multi-Objective Recommendation "
        "for QMOF Materials Discovery"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run(
        "First manuscript draft with implementation, evaluation, figures, "
        "and reproducibility appendix"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = MID_GRAY

    meta = [
        ("System", "QMOF-Rec"),
        ("Optimization contribution", "Lotus Effect Algorithm reranking"),
        ("Dataset artifact", "20,372 QMOF vector metadata records"),
        ("Evaluation", "Five query scenarios, six ranking methods, top-5 recommendations"),
        ("Draft status", "Initial manuscript draft for revision and extension"),
    ]
    add_key_value_table(doc, meta, [2200, 7160])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.add_run("Prepared for internal research development").italic = True
    doc.add_page_break()


def add_key_value_table(doc: Document, rows, widths) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, widths)
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = str(key)
        table.cell(idx, 1).text = str(value)
        set_cell_shading(table.cell(idx, 0), LIGHT_GRAY)
        for cell in table.rows[idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        table.cell(idx, 0).paragraphs[0].runs[0].font.bold = True


def add_paragraphs(doc: Document, paragraphs) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        p.add_run(item)


def add_numbered(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(6)
        p.add_run(item)


def add_heading_page(doc: Document, title: str) -> None:
    doc.add_page_break()
    doc.add_heading(title, level=1)


def add_figure(doc: Document, path: Path, caption: str, width=6.2) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption, style="CaptionCustom")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_dataframe_table(
    doc: Document,
    df: pd.DataFrame,
    columns,
    widths,
    max_rows=None,
) -> None:
    data = df[columns].copy()
    if max_rows is not None:
        data = data.head(max_rows)

    table = doc.add_table(rows=len(data) + 1, cols=len(columns))
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])

    for col_idx, column in enumerate(columns):
        cell = table.cell(0, col_idx)
        cell.text = column.replace("_", " ").title()
        set_cell_shading(cell, LIGHT_GRAY)
        cell.paragraphs[0].runs[0].font.bold = True

    for row_idx, (_, row) in enumerate(data.iterrows(), start=1):
        for col_idx, column in enumerate(columns):
            value = row[column]
            if isinstance(value, float):
                value = f"{value:.4f}"
            table.cell(row_idx, col_idx).text = str(value)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def top_method_text(aggregate: pd.DataFrame) -> dict:
    result = {}
    for metric in [
        "mean_relevance",
        "diversity",
        "ndcg_at_k",
        "hypervolume_proxy",
        "runtime_ms",
    ]:
        ascending = metric == "runtime_ms"
        row = aggregate.sort_values(metric, ascending=ascending).iloc[0]
        result[metric] = (row["method"], float(row[metric]))
    return result


def add_algorithm_box(doc: Document) -> None:
    rows = [
        ("Input", "Retrieved candidate pool C, objective weights w, population size N, iterations T, top-k K."),
        ("Initialize", "Seed population with candidate objective vectors and random normalized vectors."),
        ("Evaluate", "Map every individual to its nearest QMOF objective vector and compute LEA fitness."),
        ("Mutation", "Move each individual toward the best solution using a decaying lotus factor and random local perturbation."),
        ("Selection", "Accept mutants only when their fitness improves the current individual."),
        ("Self-cleaning", "Every tenth iteration, replace the weakest one-fifth of the population with fresh random solutions."),
        ("Output", "Return unique QMOFs ranked by LEA fitness with LEA score, rank, and optimization metadata."),
    ]
    add_key_value_table(doc, rows, [1900, 7460])


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(EVAL_DIR / "summary_metrics.csv")
    aggregate = pd.read_csv(EVAL_DIR / "aggregate_metrics.csv")
    rankings = pd.read_csv(EVAL_DIR / "top_rankings.csv")
    notes = pd.read_json(EVAL_DIR / "experiment_notes.json", typ="series")
    leaders = top_method_text(aggregate)

    doc = Document()
    style_document(doc)
    add_title_page(doc)

    doc.add_heading("Abstract", level=1)
    add_paragraphs(doc, [
        "Metal-organic frameworks provide a large design space for adsorption, catalysis, storage, and electronic applications. The size and heterogeneity of the QMOF search space make manual screening difficult, especially when a researcher must balance semantic intent, electronic suitability, density, stability, and novelty. This manuscript introduces a Lotus Effect Algorithm (LEA)-guided recommendation layer for QMOF-Rec, a retrieval-augmented materials discovery system.",
        "The proposed system retrieves a candidate pool from vector metadata and reranks candidates using a population-based optimizer inspired by lotus mutation and self-cleaning. Each candidate is represented by normalized objective scores, including semantic relevance, band-gap suitability, density suitability, porosity, and stability. LEA searches the objective space for balanced candidates while retaining diversity among top recommendations.",
        f"The first evaluation used {int(notes['materials'])} QMOF metadata records, five query scenarios, six ranking methods, and top-{int(notes['top_k'])} recommendation lists. LEA achieved mean NDCG@K of {aggregate.loc[aggregate.method == 'LEA', 'ndcg_at_k'].iloc[0]:.3f}, remained close to the highest weighted relevance baseline, and produced substantially higher diversity than weighted-sum ranking. The study also identifies an important data limitation: void fraction is absent in the current vector metadata, so porosity-aware conclusions require re-indexing with complete QMOF descriptors.",
    ])
    doc.add_heading("Keywords", level=2)
    doc.add_paragraph("QMOF; metal-organic frameworks; recommender systems; Lotus Effect Algorithm; multi-objective optimization; retrieval-augmented generation; materials discovery; explainable AI.")

    add_heading_page(doc, "1. Introduction")
    add_paragraphs(doc, [
        "Materials discovery increasingly depends on computational databases, learned representations, and interactive decision support. Metal-organic frameworks (MOFs) are especially challenging because a single application can depend on multiple coupled properties, including pore structure, density, stability, electronic band gap, topology, and chemical composition. A useful recommender must therefore avoid reducing the problem to a single static score too early.",
        "QMOF-Rec addresses this challenge by combining retrieval, property scoring, graph-aware representations, and conversational scientific assistance. The present work adds a new optimization contribution: a Lotus Effect Algorithm-guided reranking stage that treats recommendation as a multi-objective search over retrieved QMOF candidates.",
        "The core idea is to use retrieval for recall and LEA for decision quality. Retrieval produces a candidate pool related to the user query. Objective scorers transform each candidate into a normalized profile. LEA then searches this profile space, preserving candidates that balance query relevance and physicochemical suitability while periodically replacing weak solutions through a self-cleaning operation.",
    ])
    doc.add_heading("1.1 Motivation", level=2)
    add_bullets(doc, [
        "Single-score ranking is easy to implement but can hide trade-offs between relevance, density, band gap, stability, and diversity.",
        "Standard Pareto methods expose trade-offs but may produce large fronts that still require secondary ranking.",
        "A metaheuristic optimizer can search the retrieved candidate region while preserving diversity and objective balance.",
        "LEA is attractive because its self-cleaning metaphor maps naturally onto recommendation: weak candidate neighborhoods are periodically replaced to maintain exploration.",
    ])
    doc.add_heading("1.2 Contributions", level=2)
    add_numbered(doc, [
        "An end-to-end QMOF-AI platform for materials discovery, integrating retrieval, graph learning, property prediction, scientific assistance, and recommendation.",
        "A LEA-guided post-retrieval multi-objective reranking strategy for QMOF recommendation.",
        "A graph-aware recommendation extension based on GraphSAGE/GAT representations.",
        "A retrieval-augmented scientific assistant layer for evidence-grounded explanation and interaction.",
        "A reproducible evaluation framework for ranking quality, with a clear protocol for future graph-model and LLM evaluation.",
    ])

    add_heading_page(doc, "2. End-to-End QMOF-AI Platform Architecture")
    add_paragraphs(doc, [
        "This manuscript presents not only LEA-based reranking, but also a broader QMOF-AI platform for materials discovery. The platform integrates retrieval-augmented scientific assistance, graph-aware representation learning, property prediction, and multi-objective recommendation for QMOF discovery.",
        "Figure 1 is a newly constructed overview figure of the proposed platform. It is not a reused uploaded platform image. The figure was redrawn from the updated system description so that the manuscript architecture is internally consistent with the current backend, frontend, retrieval, graph, recommendation, infrastructure, and feedback-loop design.",
    ])
    add_figure(
        doc,
        ARCH_FIGURE,
        "Figure 1. Overview of the end-to-end QMOF-AI platform for materials discovery, integrating frontend interaction, backend services, retrieval-augmented scientific assistance, graph learning, property prediction, multi-objective recommendation, and user feedback.",
        width=6.5,
    )
    doc.add_heading("2.1 Module Status and Scope", level=2)
    add_paragraphs(doc, [
        "The recommendation module is the component most directly evaluated in this draft. Candidate retrieval, objective scoring, weighted baselines, TOPSIS/Pareto baselines, LEA-guided reranking, and final top-K output are implemented in the evaluation workflow.",
        "The graph-learning path is architecture-supported and partially implemented through graph construction and GraphSAGE-related components. GAT benchmarking, graph-embedding similarity, and complete graph-aware recommendation evaluation remain planned extensions. They are therefore described as platform-supported capabilities and evaluation targets rather than completed empirical results.",
        "The RAG and scientific assistant layer is implemented at the service level, but LLM faithfulness, citation quality, uncertainty calibration, and user-facing answer quality are not fully benchmarked in this draft. The manuscript defines a protocol for future LLM evaluation without fabricating results.",
    ])
    doc.add_heading("2.2 Evaluation Boundaries", level=2)
    add_bullets(doc, [
        "Fully evaluated in this draft: LEA-guided ranking quality against deterministic and multi-objective baselines.",
        "Implemented but not fully benchmarked: frontend interaction, RAG service, CIF/structure handling, graph construction, and GraphSAGE-style prediction support.",
        "Planned extensions: full GAT evaluation, graph-aware ranking ablations, LLM faithfulness evaluation, user feedback learning, and expert review.",
    ])

    add_heading_page(doc, "3. Background and Related Work")
    add_paragraphs(doc, [
        "The QMOF database provides quantum-chemical properties for a large collection of experimentally derived and hypothetical MOFs. Such datasets enable screening, supervised property prediction, and retrieval-based exploration. However, database scale alone does not solve the ranking problem: candidates must be prioritized according to application-specific objectives.",
        "Retrieval-augmented generation and vector search provide flexible access to textual and metadata representations. FAISS-style approximate or exact vector search can efficiently retrieve candidates, but retrieval distance is not identical to material suitability. A retrieved material may be semantically similar but poor in density or electronic profile.",
        "Multi-criteria decision methods such as TOPSIS and Pareto ranking are common ways to reason about trade-offs. Weighted sums are simple and interpretable, but they can over-concentrate recommendations around one narrow optimum. Pareto approaches are more faithful to trade-offs but still need ranking inside fronts. LEA is evaluated here as a search-based alternative that can combine weighted objective fitness, balance, and diversity.",
    ])
    doc.add_heading("2.1 Lotus Effect Algorithm", level=2)
    add_paragraphs(doc, [
        "The LEA paper introduces a nature-inspired evolutionary optimizer based on lotus pollination, dragonfly-inspired exploration, and self-cleaning behavior. In this manuscript, the implemented version follows the simplified LEA code variant supplied with the project: candidates move toward the current best solution using a decaying mutation factor, while the weakest portion of the population is periodically replaced.",
        "This adaptation is intentionally recommendation-oriented. Rather than generating new crystal structures, LEA searches a normalized objective space and maps each individual back to the nearest retrieved QMOF candidate. This preserves database validity: every returned recommendation corresponds to an existing material record.",
    ])

    add_heading_page(doc, "4. System Architecture")
    add_paragraphs(doc, [
        "The system is organized as a modular web application. The backend exposes FastAPI routes for recommendation, chat, material prediction, and structure retrieval. The frontend is a React workspace that displays query controls, dynamic weights, recommendations, explanations, and 3D structure viewing.",
        "The recommendation pipeline begins with user intent. Dynamic weights are generated from the query, candidates are retrieved from the FAISS vector store, physicochemical scores are computed, and a final optimizer reranks candidates. Before this work, the optimizer stage was a simple final-score sort. The new implementation replaces that placeholder with LEA-based optimization metadata and rankings.",
    ])
    add_key_value_table(doc, [
        ("Frontend", "React/Vite workspace with recommendation panel, metric display, chat, and material cards."),
        ("Backend API", "FastAPI routes for recommendation, chat, material prediction, and structure retrieval."),
        ("Retrieval layer", "Sentence-transformer embeddings and FAISS vector index over QMOF metadata."),
        ("Scoring layer", "Semantic, band-gap, density, porosity, stability, and similarity scores."),
        ("Optimization layer", "LEA reranking over normalized objective vectors."),
        ("Evaluation layer", "Standalone benchmark script producing CSV metrics and manuscript figures."),
    ], [1800, 7560])
    doc.add_heading("3.1 Recommendation Flow", level=2)
    add_numbered(doc, [
        "Receive a natural-language materials discovery query.",
        "Generate query-adaptive objective weights.",
        "Retrieve a candidate pool from the vector store.",
        "Compute normalized objective scores for each candidate.",
        "Run LEA over the candidate objective matrix.",
        "Return ranked QMOFs with LEA score, LEA rank, final score, and explanations.",
    ])

    add_heading_page(doc, "5. LEA-Guided Recommendation Method")
    add_paragraphs(doc, [
        "The recommendation adaptation of LEA treats each individual as a point in normalized objective space. The dimensions are semantic relevance, band-gap suitability, density suitability, porosity, and stability. Since candidates are fixed database records, continuous individuals are mapped to the nearest candidate objective vector during evaluation.",
        "Fitness combines three terms. The first is query-weighted suitability. The second is objective balance, computed from the weakest objective. The third is diversity relative to already selected high-quality candidates. This design encourages LEA to find candidates that are relevant, balanced, and not redundant.",
    ])
    add_algorithm_box(doc)
    doc.add_heading("4.1 Fitness Function", level=2)
    add_paragraphs(doc, [
        "Let x denote an LEA individual and c(x) the nearest retrieved QMOF candidate in objective space. Let s(c) be the vector of normalized objective scores and w be the query-adaptive weight vector. The implemented fitness is:",
        "fitness(x) = dot(s(c), w) + 0.12 * min(s(c)) + 0.08 * diversity(s(c), S)",
        "Here S is the set of currently selected high-quality candidate vectors. The balance term rewards candidates that avoid a severe weakness, while the diversity term discourages collapse into near-duplicate recommendations.",
    ])
    doc.add_heading("4.2 Implementation Notes", level=2)
    add_bullets(doc, [
        "Population size is 30 and maximum iterations are 60 in the current evaluation.",
        "The optimizer is deterministic under a fixed random seed.",
        "The final API response includes LEA metadata so frontend users can see that optimization was applied.",
        "The implementation is extensible: the full dragonfly separation, alignment, cohesion, food, and enemy terms from the LEA paper can be added later.",
    ])

    add_heading_page(doc, "6. Data and Property Coverage")
    add_paragraphs(doc, [
        f"The current experiment used {int(notes['materials'])} records from backend/vector_db/metadata.json. This is the vector metadata available in the workspace. The full QMOF CSV was not present, so the evaluation did not assume unavailable descriptor columns.",
        "The available metadata contains density for all materials, band gap for approximately 53.1 percent of materials, and no populated void-fraction values. This directly affects porosity and balance metrics. A scientifically stronger follow-up experiment should rebuild the vector store from the full QMOF table with void fraction, pore-limiting diameter, largest cavity diameter, surface area, synthesized flag, topology, and source metadata.",
    ])
    add_figure(doc, EVAL_DIR / "figure_1_data_coverage.png", "Figure 2. Property coverage in the current vector metadata. Void fraction is absent, so porosity-aware evaluation is limited in this first run.")
    doc.add_heading("5.1 Consequence for Interpretation", level=2)
    add_paragraphs(doc, [
        "Because porosity is unavailable, all porosity scores are zero in this benchmark. Balance and hypervolume proxy values are therefore depressed. The results should be read as an implementation and ranking-behavior evaluation under incomplete metadata, not as a final materials-science validation of porosity-aware recommendation.",
        "This limitation is useful rather than merely inconvenient: it exposes the difference between a working recommendation architecture and a publication-ready descriptor pipeline. The manuscript therefore reports LEA behavior honestly and identifies data restoration as the next experimental priority.",
    ])

    add_heading_page(doc, "7. Experimental Design")
    add_paragraphs(doc, [
        "The benchmark evaluates ranking quality across five query scenarios selected to represent common MOF discovery intents: CO2 adsorption, photocatalysis, lightweight gas storage, balanced materials discovery, and wide-band-gap insulating frameworks.",
        "Each query defines a deterministic objective weight vector. A candidate pool of 100 materials is constructed using a deterministic lexical and property proxy. This design avoids external API calls or model downloads, making the experiment reproducible in the current workspace.",
    ])
    query_rows = [
        ("CO2 adsorption", "Porosity and stability emphasized; semantic and density secondary."),
        ("Photocatalysis", "Band-gap suitability emphasized with stability support."),
        ("Lightweight storage", "Density and porosity emphasized."),
        ("Balanced discovery", "All non-porosity objectives weighted moderately."),
        ("Insulating frameworks", "Wide band gap and density emphasized."),
    ]
    add_key_value_table(doc, query_rows, [2500, 6860])
    doc.add_heading("6.1 Baselines", level=2)
    add_bullets(doc, [
        "SemanticOnly ranks by semantic relevance proxy.",
        "WeightedSum ranks by the query-weighted objective average.",
        "TOPSIS ranks by closeness to the ideal objective vector.",
        "ParetoCrowding applies non-dominated sorting with crowding and weighted tie-breaking.",
        "Random provides a stochastic lower reference.",
        "LEA applies population-based lotus mutation and self-cleaning reranking.",
    ])
    doc.add_heading("6.2 Metrics", level=2)
    add_bullets(doc, [
        "Mean relevance: average query-weighted objective score among top-k materials.",
        "NDCG@K: ranking quality relative to the ideal weighted ordering within the candidate pool.",
        "Diversity: mean pairwise distance among top-k objective vectors.",
        "Balance: average minimum objective score across recommended candidates.",
        "Hypervolume proxy: product-based objective coverage proxy, reported cautiously because porosity is absent.",
        "Runtime: average wall-clock ranking time in milliseconds.",
    ])

    add_heading_page(doc, "8. Aggregate Results")
    add_paragraphs(doc, [
        f"WeightedSum achieved the highest mean relevance ({leaders['mean_relevance'][1]:.3f}), as expected because the metric is directly aligned with weighted scoring. LEA achieved mean relevance of {aggregate.loc[aggregate.method == 'LEA', 'mean_relevance'].iloc[0]:.3f}, very close to the weighted optimum.",
        f"LEA achieved NDCG@K of {aggregate.loc[aggregate.method == 'LEA', 'ndcg_at_k'].iloc[0]:.3f}. Its main advantage is diversity: LEA produced mean diversity of {aggregate.loc[aggregate.method == 'LEA', 'diversity'].iloc[0]:.3f}, compared with {aggregate.loc[aggregate.method == 'WeightedSum', 'diversity'].iloc[0]:.3f} for WeightedSum. This indicates that LEA preserves alternative candidates while remaining close to the relevance optimum.",
    ])
    metric_cols = ["method", "mean_relevance", "diversity", "ndcg_at_k", "hypervolume_proxy", "runtime_ms"]
    add_dataframe_table(doc, aggregate.sort_values("mean_relevance", ascending=False), metric_cols, [1700, 1500, 1300, 1300, 1700, 1360])
    add_figure(doc, EVAL_DIR / "figure_2_metric_comparison.png", "Figure 3. Aggregate comparison across ranking methods. LEA is close to WeightedSum in relevance and NDCG while producing much higher diversity.")

    add_heading_page(doc, "9. Runtime and Computational Cost")
    add_paragraphs(doc, [
        "The optimization methods differ substantially in runtime. WeightedSum, SemanticOnly, TOPSIS, and Random are sub-millisecond to near-sub-millisecond in this benchmark. ParetoCrowding is slower because of pairwise dominance comparisons. LEA is the slowest method because it performs iterative population evaluation.",
        "This runtime is acceptable for interactive recommendation when candidate pools are modest, but it should be optimized for larger-scale search. Potential improvements include vectorized fitness caching, adaptive early stopping, lower population size for exploratory UI queries, and full-size LEA only when the user requests optimization-grade recommendations.",
    ])
    add_figure(doc, EVAL_DIR / "figure_3_runtime.png", "Figure 4. Average runtime by ranking method. LEA is slower than deterministic baselines but remains practical for reranking a retrieved pool.")

    add_heading_page(doc, "10. LEA Convergence")
    add_paragraphs(doc, [
        "LEA convergence curves show rapid early improvement followed by stable plateaus. This behavior is expected because the candidate pool is discrete and the population is seeded with candidate objective vectors. Once LEA identifies high-fitness candidates, later iterations primarily maintain or slightly refine the best solution.",
        "The convergence behavior suggests that the current 60 iterations may be more than necessary for some queries. A future adaptive version can stop when best fitness does not improve for a fixed patience window.",
    ])
    add_figure(doc, EVAL_DIR / "figure_4_lea_convergence.png", "Figure 5. Best LEA fitness over 60 iterations for five query scenarios. Most gains occur early.")

    add_heading_page(doc, "11. Objective Profiles")
    add_paragraphs(doc, [
        "The objective radar plot compares the mean profile of LEA and WeightedSum recommendations. LEA improves band-gap suitability relative to WeightedSum while maintaining similar semantic, density, and stability profiles. Porosity remains zero for both because void fraction is missing in the current metadata.",
        "This profile supports a nuanced interpretation: LEA does not simply maximize the same weighted score. It can shift the recommendation set toward a richer objective profile, especially when the fitness includes balance and diversity terms.",
    ])
    add_figure(doc, EVAL_DIR / "figure_5_objective_radar.png", "Figure 6. Mean objective profile for LEA and WeightedSum top recommendations.")

    add_heading_page(doc, "12. Query-Level Results")
    add_paragraphs(doc, [
        "Query-level results show that LEA often matches the best deterministic methods in NDCG while producing broader candidate sets. The largest diversity gain appears in the lightweight storage and insulating-framework scenarios, where LEA and ParetoCrowding identify more varied objective profiles.",
        "The CO2 adsorption scenario is the most constrained by missing porosity metadata. In a complete QMOF descriptor table, this query should benefit most from pore descriptors such as void fraction, pore-limiting diameter, largest cavity diameter, accessible surface area, and density.",
    ])
    qcols = ["query_id", "method", "mean_relevance", "diversity", "ndcg_at_k", "runtime_ms"]
    add_dataframe_table(doc, summary.sort_values(["query_id", "method"]), qcols, [2300, 1700, 1400, 1300, 1300, 1160], max_rows=30)

    add_heading_page(doc, "13. Top Recommendation Examples")
    add_paragraphs(doc, [
        "The table below lists the first LEA recommendation for each query scenario. These examples are not presented as final validated materials; they are evidence that the pipeline returns concrete QMOF IDs with traceable objective scores and can be inspected further through structure retrieval and domain validation.",
    ])
    lea_first = rankings[(rankings.method == "LEA") & (rankings["rank"] == 1)].copy()
    rcols = ["query_id", "qmof_id", "formula", "band_gap", "density", "lea_score"]
    add_dataframe_table(doc, lea_first, rcols, [2300, 1600, 2500, 1200, 1200, 1160])
    doc.add_heading("12.1 Interpretation Workflow", level=2)
    add_numbered(doc, [
        "Inspect LEA top candidates and compare with WeightedSum alternatives.",
        "Open each material card in the frontend to review scores and explanations.",
        "Load CIF structure when available to inspect geometry.",
        "Run graph-based prediction or descriptor extraction for candidate validation.",
        "Export candidates for expert review or DFT/simulation follow-up.",
    ])

    add_heading_page(doc, "14. Discussion")
    add_paragraphs(doc, [
        "The first implementation demonstrates that LEA can be integrated into a retrieval-based QMOF recommender without changing the external API contract dramatically. The frontend can display LEA rank and score, while the backend can still return familiar material properties and explanations.",
        "The primary scientific value of LEA is not raw weighted relevance. WeightedSum will usually win a metric that is itself a weighted sum. LEA is valuable when the desired recommendation list should include diverse, balanced, and application-plausible candidates. This is especially important in early-stage materials discovery, where researchers often prefer a portfolio of candidate families rather than five near-duplicates.",
        "The current LEA adaptation also creates an interpretable bridge between metaheuristic optimization and recommender systems. The self-cleaning step periodically removes weak solution neighborhoods, which helps avoid stagnation and encourages exploration of alternative candidate regions.",
    ])
    doc.add_heading("13.1 What the Current Results Support", level=2)
    add_bullets(doc, [
        "LEA can be implemented as a post-retrieval reranker for QMOF candidates.",
        "LEA remains close to top weighted relevance and NDCG performance.",
        "LEA improves recommendation diversity over WeightedSum.",
        "The evaluation harness can compare optimization algorithms reproducibly.",
    ])
    doc.add_heading("13.2 What the Current Results Do Not Yet Prove", level=2)
    add_bullets(doc, [
        "They do not prove superior physical performance for CO2 adsorption because porosity metadata is missing.",
        "They do not validate synthesis feasibility experimentally.",
        "They do not compare against graph embeddings yet.",
        "They do not use the full formal dragonfly exploration equations from the LEA paper.",
    ])

    add_heading_page(doc, "15. Limitations and Threats to Validity")
    add_paragraphs(doc, [
        "The most important limitation is descriptor incompleteness. Void fraction is absent in the current vector metadata, so porosity scores are zero. This affects balance and hypervolume proxy metrics and weakens claims for adsorption-oriented tasks.",
        "The second limitation is the deterministic semantic proxy used for evaluation. The API recommender can use embeddings, but the benchmark avoids external model downloads for reproducibility. A final paper should repeat evaluation with the actual retrieval model or a frozen local embedding model.",
        "Third, stability is represented by a heuristic proxy. A stronger system should use QMOF source labels, synthesized flags, decomposition or formation-energy proxies when available, and chemistry-aware risk scoring.",
        "Fourth, the current LEA implementation is the simplified version supplied in code. The full LEA paper includes dragonfly-inspired exploration terms, local exploitation, and water-drop self-cleaning modeling. Implementing those terms would strengthen methodological fidelity.",
    ])

    add_heading_page(doc, "16. Future Work")
    add_numbered(doc, [
        "Rebuild the vector metadata from the full QMOF CSV with pore descriptors, topology, synthesized flag, and source metadata.",
        "Add graph embeddings from GraphSAGE or a topology-aware GNN as a structural objective.",
        "Implement the full LEA equations, including separation, alignment, cohesion, food attraction, enemy distraction, exploitation radius, and local water-drop search.",
        "Evaluate against additional metaheuristic baselines such as PSO, genetic algorithm, random search, Bayesian optimization, and NSGA-II.",
        "Add statistical testing across repeated seeds and larger query suites.",
        "Create an expert-review protocol where materials scientists judge top recommendations for plausibility and novelty.",
    ])

    add_heading_page(doc, "17. Conclusion")
    add_paragraphs(doc, [
        "This manuscript presents the first integrated draft of a LEA-guided multi-objective recommendation layer for QMOF-Rec. The system reranks retrieved QMOF candidates using a population-based optimizer that balances query relevance, property suitability, stability, and diversity.",
        "The first evaluation shows that LEA is competitive with deterministic ranking methods in relevance and NDCG while improving diversity over weighted-sum ranking. This makes LEA a promising novelty contribution for a materials recommendation paper, provided the next experimental phase restores complete QMOF descriptors and expands validation.",
        "The immediate manuscript path is clear: complete the descriptor pipeline, rerun evaluation with porosity and graph features, compare against additional optimizers, and revise the current draft into a submission-ready paper.",
    ])

    add_heading_page(doc, "References")
    refs = [
        "Dalirinia, E., Jalali, M., Yaghoobi, M., and Tabatabaee, H. Lotus effect optimization algorithm (LEA): a lotus nature-inspired algorithm for engineering design optimization. The Journal of Supercomputing, 80, 761-799, 2024. https://doi.org/10.1007/s11227-023-05513-8",
        "Rosen, A. S. et al. Machine learning the quantum-chemical properties of metal-organic frameworks for accelerated materials discovery. Matter, 2021.",
        "Johnson, J., Douze, M., and Jegou, H. Billion-scale similarity search with GPUs. IEEE Transactions on Big Data, 2019.",
        "Hwang, C. L. and Yoon, K. Multiple Attribute Decision Making: Methods and Applications. Springer, 1981.",
        "Deb, K., Pratap, A., Agarwal, S., and Meyarivan, T. A fast and elitist multiobjective genetic algorithm: NSGA-II. IEEE Transactions on Evolutionary Computation, 2002.",
        "Hamilton, W. L., Ying, R., and Leskovec, J. Inductive representation learning on large graphs. NeurIPS, 2017.",
        "Lewis, P. et al. Retrieval-augmented generation for knowledge-intensive NLP tasks. NeurIPS, 2020.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="BodySmall")

    add_heading_page(doc, "Appendix A. Reproducibility Checklist")
    add_key_value_table(doc, [
        ("Repository", str(ROOT.parent)),
        ("Evaluation script", "backend/scripts/evaluate_lea_recommender.py"),
        ("Evaluation output", "reports/lea_evaluation"),
        ("Manuscript builder", "reports/manuscript/build_manuscript.py"),
        ("Candidate pool size", str(int(notes["candidate_pool_size"]))),
        ("Top-k", str(int(notes["top_k"]))),
        ("Methods", ", ".join(notes["methods"])),
    ], [2500, 6860])
    doc.add_heading("A.1 Commands", level=2)
    commands = [
        "MPLCONFIGDIR=/private/tmp/qmof-matplotlib PYTHONPATH=backend python3 backend/scripts/evaluate_lea_recommender.py --candidate-pool-size 100 --top-k 5 --out-dir reports/lea_evaluation",
        "/Users/mehrdadjalali/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3 reports/manuscript/build_manuscript.py",
    ]
    for command in commands:
        p = doc.add_paragraph()
        r = p.add_run(command)
        r.font.name = "Courier New"
        r.font.size = Pt(8)

    add_heading_page(doc, "Appendix B. Evaluation Tables")
    add_paragraphs(doc, [
        "This appendix contains compact copies of the aggregate and query-level evaluation tables. The full machine-readable outputs are stored as CSV files in reports/lea_evaluation.",
    ])
    add_dataframe_table(doc, aggregate.sort_values("method"), metric_cols, [1700, 1500, 1300, 1300, 1700, 1360])
    doc.add_heading("B.1 LEA Top-Five Rows", level=2)
    add_dataframe_table(doc, rankings[rankings.method == "LEA"].head(25), ["query_id", "rank", "qmof_id", "formula", "band_gap", "density", "lea_score"], [2100, 700, 1500, 2300, 1000, 1000, 760], max_rows=25)

    add_heading_page(doc, "Appendix C. Data Restoration Plan")
    add_paragraphs(doc, [
        "The current manuscript draft should be treated as a first experimental report. To make it submission-ready, the QMOF descriptor table should be restored and the vector index rebuilt. The target metadata schema should include at least the following fields: qmof_id, formula, reduced formula, topology, source, synthesized flag, density, void fraction, pore-limiting diameter, largest cavity diameter, surface area, band gap, and structural embedding identifiers.",
        "After rebuilding the index, the benchmark should be rerun with the same scripts. Porosity-aware metrics should then become meaningful, balance should no longer be forced to zero, and adsorption-oriented query scenarios can be interpreted with stronger materials-science validity.",
    ])
    add_bullets(doc, [
        "Add descriptor ingestion tests to catch missing porosity columns.",
        "Report coverage before every evaluation run.",
        "Store evaluation configuration as JSON alongside every result table.",
        "Use repeated seeds for LEA and stochastic baselines.",
        "Include statistical tests and confidence intervals in the revised manuscript.",
    ])

    add_heading_page(doc, "Appendix D. Detailed Algorithmic Rationale")
    add_paragraphs(doc, [
        "The LEA reranking layer is designed as a post-retrieval optimizer rather than as a replacement for retrieval. This distinction is important. Retrieval is responsible for recall: it should gather a pool of plausible candidates from a large database. Optimization is responsible for portfolio construction: it should choose a small set of candidates that reflects the user's objective priorities while retaining enough diversity for scientific exploration.",
        "The optimizer therefore operates on normalized objective vectors rather than raw material records. This makes the method independent of any one descriptor scale and allows new objectives to be added later. For example, a graph-embedding similarity score, synthesis-risk score, or uncertainty score can become another coordinate in the objective vector without changing the retrieval layer.",
        "Mapping continuous LEA individuals to the nearest candidate is a pragmatic design choice. A pure continuous optimizer could generate an idealized vector that does not correspond to any known QMOF. In a recommendation system, that would be misleading unless the system also included a valid generative model. Nearest-candidate mapping ensures that every recommendation is traceable to an existing QMOF identifier.",
        "The balance term is intentionally small but meaningful. A high weighted score can still hide a severe weakness in one dimension. For materials discovery, such weaknesses matter: a candidate with strong semantic relevance and band-gap suitability may still be undesirable if density or stability is poor. The minimum-objective balance term rewards candidates that avoid such severe weaknesses.",
        "The diversity term is also intentionally small. It should not dominate relevance, but it should discourage the top-k list from collapsing into closely related candidates. In early-stage screening, diversity is often valuable because researchers may prefer several chemically or structurally distinct families for follow-up validation.",
    ])
    doc.add_heading("D.1 Full LEA Fidelity Roadmap", level=2)
    add_paragraphs(doc, [
        "The current implementation follows the simplified LEA code supplied for this project. A full paper-grade extension should implement the original LEA exploration model more faithfully. In particular, the dragonfly-inspired terms can be incorporated as population-level operators over the objective space. Separation would push individuals away from crowded neighborhoods. Alignment would coordinate movement among nearby promising individuals. Cohesion would pull individuals toward local candidate clusters. Food attraction would use the best candidate or best front member, and enemy distraction would push the population away from low-fitness regions.",
        "The exploitation radius from the LEA paper can also be added as a decaying local-search radius. Early iterations would permit broad movement around the current best solution, while later iterations would restrict movement to small refinements. This would make convergence behavior more faithful to the lotus pollination metaphor and may reduce the number of wasted late iterations.",
        "Finally, the self-cleaning model can be extended beyond random replacement. Weak individuals could be redirected toward higher-capacity candidate neighborhoods, where capacity is derived from fitness, diversity contribution, or non-dominated front membership. This would connect the implemented recommendation method more directly to the water-drop and pit-capacity formulation in the original LEA paper.",
    ])

    add_heading_page(doc, "Appendix E. Full-Scale Evaluation Plan")
    add_paragraphs(doc, [
        "The present benchmark is deliberately reproducible in the current workspace, but a final manuscript should include a larger evaluation. The recommended full-scale experiment has three phases. First, rebuild metadata from the complete QMOF source table. Second, run query-driven ranking experiments with repeated random seeds. Third, validate recommendation quality using both automatic metrics and expert review.",
        "The first phase should compute property coverage before any ranking experiment. Coverage should be reported for density, band gap, void fraction, pore-limiting diameter, largest cavity diameter, surface area, topology, synthesized flag, source database, and any graph-derived descriptors. This step prevents hidden missingness from distorting results.",
        "The second phase should use a broader query suite. Queries should cover adsorption, catalysis, electronic materials, low-density frameworks, high-stability frameworks, and general discovery. Each query should define expected objective priorities before running the algorithms. This avoids retrofitting weights after seeing results.",
        "The third phase should evaluate scientific usefulness. Automatic metrics are necessary but insufficient. Expert review can ask whether recommended candidates are plausible, diverse, non-obvious, and worth follow-up. A small blinded expert study would substantially strengthen the manuscript.",
    ])
    doc.add_heading("E.1 Recommended Additional Baselines", level=2)
    add_bullets(doc, [
        "NSGA-II for evolutionary multi-objective ranking.",
        "Particle swarm optimization using the older project code as a baseline.",
        "Genetic algorithm ranking over normalized objective vectors.",
        "Bayesian optimization over candidate objective space.",
        "Maximal marginal relevance for diversity-aware retrieval.",
        "Graph-only similarity ranking using GraphSAGE embeddings.",
        "Hybrid graph and text retrieval without LEA.",
    ])
    doc.add_heading("E.2 Statistical Reporting", level=2)
    add_paragraphs(doc, [
        "For stochastic algorithms, each experiment should be repeated with multiple seeds. The final manuscript should report mean, standard deviation, and confidence intervals for NDCG@K, diversity, relevance, hypervolume, and runtime. Pairwise comparisons can use non-parametric tests when metric distributions are not normal.",
        "Runtime should be reported separately from recommendation quality. A slower optimizer can be acceptable if it provides better candidate portfolios, but the trade-off should be explicit. A practical deployment can use a fast weighted baseline for instant preview and LEA for optimized recommendation mode.",
    ])

    add_heading_page(doc, "Appendix F. Figure Interpretation Notes")
    add_paragraphs(doc, [
        "Figure 1 should appear early in the manuscript because it frames the validity of every downstream result. It shows that the current metadata is sufficient for density-aware and partial band-gap-aware experiments, but not yet sufficient for pore-aware claims. This figure prevents overstatement and gives reviewers confidence that data limitations were inspected.",
        "Figure 2 summarizes the main quantitative comparison. The key interpretation is not that LEA beats WeightedSum on mean relevance. It does not, and it should not necessarily do so when relevance is defined as a weighted sum. The key interpretation is that LEA remains near the top in NDCG and relevance while increasing diversity.",
        "Figure 3 should be discussed candidly. LEA is slower than simple deterministic methods. The manuscript should treat this as an engineering trade-off rather than hiding it. A recommender can justify the extra runtime when producing a small, optimized shortlist for research follow-up.",
        "Figure 4 supports a practical early-stopping improvement. Most LEA runs plateau quickly, suggesting that future versions can reduce iteration count or stop after a patience window. This can reduce runtime while preserving recommendation quality.",
        "Figure 5 is useful for explaining objective behavior to readers. It shows that LEA changes the profile of top recommendations rather than simply reordering by the same final score. The zero porosity axis should be explicitly tied back to missing void-fraction metadata.",
    ])
    doc.add_heading("F.1 Suggested Caption Style for Submission", level=2)
    add_paragraphs(doc, [
        "For a journal submission, captions should be expanded beyond the compact captions in this first draft. Each caption should state the data source, methods compared, metric definition, and main interpretation. Reviewers often read figures before the method section, so self-contained captions will make the contribution easier to understand.",
        "The figure set should also be supplemented with a system architecture diagram and a flow diagram of the LEA reranking process. These two diagrams are not generated yet in this first draft, but they would improve the manuscript's narrative structure.",
    ])

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
